#!/usr/bin/env python3
"""
latex_to_docx_all_v2.py  —  with brace hotspot diagnostics and optional auto-fix

Adds:
- Block-aware stripping of macro definitions (resume helpers and structural redefs).
- Brace hotspot reporting: show exact line/column and context for any unmatched '{' left.
- Optional --auto-fix-braces: if brace imbalance > 0, insert that many '}' just before \end{document}.
  (Last-resort patch; prefer to fix the true source when possible.)
- Everything else from v1: guard glyphtounicode, normalize resume macros, balance envs, ensure list items,
  create reference.docx, and optionally run pandoc.
"""

import argparse, os, re, shutil, subprocess, sys, tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ---------------------- reference.docx generation ----------------------
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CMU = {"serif": "CMU Serif", "sans": "CMU Sans Serif", "mono": "CMU Typewriter Text"}
LM  = {"serif": "Latin Modern Roman", "sans": "Latin Modern Sans", "mono": "Latin Modern Mono"}

def _set_style_font(doc: Document, style: str, name: str, size_pt: Optional[float] = None,
                    bold: Optional[bool] = None, italic: Optional[bool] = None) -> None:
    try:
        st = doc.styles[style]
    except KeyError:
        return
    f = st.font
    f.name = name
    if size_pt is not None:
        f.size = Pt(size_pt)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic

def make_reference_docx(fonts: Dict[str, str], base_size: float = 11.0) -> str:
    doc = Document()
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Reference Styles"); r.bold = True
    doc.add_paragraph("Normal body")

    _set_style_font(doc, "Normal",    fonts["serif"], base_size)
    _set_style_font(doc, "Heading 1", fonts["serif"], base_size+5, bold=True)
    _set_style_font(doc, "Heading 2", fonts["serif"], base_size+3, bold=True)
    _set_style_font(doc, "Heading 3", fonts["serif"], base_size+2, bold=True)
    _set_style_font(doc, "Block Text", fonts["serif"], base_size, italic=True)
    try:
        doc.styles["Verbatim Char"].font.name = fonts["mono"]
    except KeyError:
        pass
    _set_style_font(doc, "Source Code", fonts["mono"], base_size)

    tmp = tempfile.NamedTemporaryFile(prefix="ref_", suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name

# ---------------------- normalization helpers ----------------------

RE_ENV_BEGIN = re.compile(r'\\begin\{([A-Za-z*]+)\}')
RE_ENV_END   = re.compile(r'\\end\{([A-Za-z*]+)\}')

BALANCE_ENVS = {
    "itemize","enumerate","description","tabular","tabular*",
    "center","flushleft","flushright","quote","quotation","list"
}

def guard_glyphtounicode(text: str) -> str:
    return re.sub(
        r'^\s*\\input\{glyphtounicode(\.tex)?\}\s*$',
        r'\\InputIfFileExists{glyphtounicode.tex}{}{} % guarded for Pandoc',
        text, flags=re.M
    )

# -------- Block-aware stripping of macro definitions --------

MACRO_NAME_PATTERNS = [
    r'\\resume[A-Za-z]+',
    r'\\begin\{itemize\}', r'\\end\{itemize\}',
    r'\\labelitemi\+?',
]

RE_CMD_START = re.compile(r'\\(newcommand|renewcommand|providecommand|def)\b')

def _should_strip_definition(header_text: str) -> bool:
    return any(re.search(p, header_text) for p in MACRO_NAME_PATTERNS)

def _skip_balanced(src: str, idx: int, open_ch: str, close_ch: str) -> int:
    if idx >= len(src) or src[idx] != open_ch:
        return idx
    depth = 0
    j = idx
    while j < len(src):
        c = src[j]
        if c == open_ch:
            depth += 1
        elif c == close_ch:
            depth -= 1
            if depth == 0:
                return j + 1
        j += 1
    return j

def _strip_one_definition(src: str, start_idx: int) -> Tuple[str, int]:
    i = start_idx
    m = RE_CMD_START.match(src, i)
    if not m:
        return src, start_idx + 1
    i = m.end()

    header_window = src[start_idx: min(len(src), i + 400)]
    if not _should_strip_definition(header_window):
        return src, i

    # Optional {name} or \def\name... forms
    while i < len(src) and src[i].isspace():
        i += 1
    if i < len(src) and src[i] == '{':
        i = _skip_balanced(src, i, '{', '}')
    else:
        # \def\name#1#2...
        if i < len(src) and src[i] == '\\':
            i += 1
            while i < len(src) and (src[i].isalpha() or src[i] == '@'):
                i += 1
            while i < len(src) and src[i] in ' #0123456789':
                i += 1

    # optional [..]
    while i < len(src) and src[i].isspace():
        i += 1
    if i < len(src) and src[i] == '[':
        i = _skip_balanced(src, i, '[', ']')

    # BODY { ... }
    while i < len(src) and src[i].isspace():
        i += 1
    if i >= len(src):
        j = src.find('\n', start_idx)
        end_idx = (j + 1) if j != -1 else len(src)
        return src[:start_idx] + '% (stripped macro definition)\n' + src[end_idx:], start_idx + 1

    if src[i] != '{':
        j = src.find('{', i)
        if j == -1:
            j = src.find('\n', i)
            end_idx = (j + 1) if j != -1 else len(src)
            return src[:start_idx] + '% (stripped macro definition)\n' + src[end_idx:], start_idx + 1
        i = j

    body_end = _skip_balanced(src, i, '{', '}')
    end_idx = body_end
    return src[:start_idx] + '% (stripped macro definition)\n' + src[end_idx:], start_idx + 1

def strip_macro_definitions(text: str) -> str:
    i = 0
    src = text
    while True:
        m = RE_CMD_START.search(src, i)
        if not m:
            break
        cmd_start = m.start()
        header_window = src[cmd_start: min(len(src), cmd_start + 500)]
        if _should_strip_definition(header_window):
            src, i = _strip_one_definition(src, cmd_start)
        else:
            i = m.end()
    return src

# -------- Replace invocations of resume macros --------

def replace_invocations(text: str) -> str:
    lines_out = []
    itemize_depth = 0
    for raw in text.splitlines():
        ln = raw
        ln = ln.replace(r'\resumeSubHeadingListStart', r'\begin{itemize}')
        ln = ln.replace(r'\resumeSubHeadingListEnd',   r'\end{itemize}')
        ln = ln.replace(r'\resumeItemListStart',       r'\begin{itemize}')
        ln = ln.replace(r'\resumeItemListEnd',         r'\end{itemize}')

        for m in RE_ENV_BEGIN.finditer(ln):
            if m.group(1) == 'itemize':
                itemize_depth += 1
        for m in RE_ENV_END.finditer(ln):
            if m.group(1) == 'itemize' and itemize_depth > 0:
                itemize_depth -= 1

        def _repl_item(m):
            nonlocal itemize_depth
            text_item = m.group(1)
            if itemize_depth == 0:
                itemize_depth += 1
                return r'\begin{itemize}' + '\n' + r'\item ' + text_item
            return r'\item ' + text_item

        ln = re.sub(r'\\resumeItem\{([^}]*)\}', _repl_item, ln)
        ln = re.sub(r'\\resumeSubItem\{([^}]*)\}', _repl_item, ln)

        lines_out.append(ln)

    text = '\n'.join(lines_out) + '\n'

    def subheading_repl(m):
        A,B,C,D = m.group(1), m.group(2), m.group(3), m.group(4)
        return (r'\item \textbf{' + A + r'}\hfill\textit{' + B + r'}\\' +
                r'\textit{' + C + r'}\hfill\textit{' + D + r'}')
    text = re.sub(
        r'\\resumeSubheading\{(.*?)\}\{(.*?)\}\{(.*?)\}\{(.*?)\}',
        subheading_repl, text, flags=re.S
    )
    return text

# -------- Balance envs & ensure list items --------

def balance_envs(text: str) -> Tuple[str, Dict[str, int]]:
    lines = text.splitlines()
    stack: List[Tuple[str,int]] = []
    unmatched_end: List[Tuple[str,int]] = []

    def is_commented(ln: str) -> bool:
        return ln.lstrip().startswith('%')

    for i, ln in enumerate(lines, 1):
        if is_commented(ln):
            continue
        for m in RE_ENV_BEGIN.finditer(ln):
            env = m.group(1)
            if env in BALANCE_ENVS:
                stack.append((env, i))
        for m in RE_ENV_END.finditer(ln):
            env = m.group(1)
            if env in BALANCE_ENVS:
                if stack and stack[-1][0] == env:
                    stack.pop()
                else:
                    unmatched_end.append((env, i))

    end_doc_idx = None
    for idx, ln in enumerate(lines):
        if re.match(r'^\s*\\end\{document\}\s*$', ln):
            end_doc_idx = idx
            break
    if end_doc_idx is None:
        end_doc_idx = len(lines)

    inserts = []
    for env, _line_no in reversed(stack):
        inserts.append(f'\\end{{{env}}}')
    if inserts:
        lines[end_doc_idx:end_doc_idx] = inserts

    stats = {"inserted_missing_ends": len(inserts), "unmatched_end_count": len(unmatched_end)}
    return '\n'.join(lines) + '\n', stats

def ensure_list_items(text: str) -> str:
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        if re.search(r'\\begin\{(itemize|enumerate)\}', lines[i]):
            j = i + 1
            found = False
            while j < len(lines) and not re.search(r'\\end\{(itemize|enumerate)\}', lines[j]):
                if re.search(r'^\s*\\item\b', lines[j]):
                    found = True
                    break
                j += 1
            if not found:
                lines.insert(i+1, r'\item')
                i += 1
        i += 1
    return '\n'.join(lines) + '\n'

# -------- Brace analysis & optional fix --------

def find_unmatched_open_braces(text: str) -> List[Tuple[int,int,str]]:
    """
    Return a list of (line, col, context_line) for every '{' that was never closed.
    Ignores comments (after %) when scanning.
    """
    stack = []
    for line_no, raw in enumerate(text.splitlines(), 1):
        line = raw.split('%', 1)[0]  # strip comments
        col = 0
        while col < len(line):
            ch = line[col]
            if ch == '{':
                stack.append((line_no, col+1, raw))
            elif ch == '}':
                if stack:
                    stack.pop()
                else:
                    # stray closing brace: we ignore here; env balancing usually handles \end mismatches
                    pass
            col += 1
    return stack  # any left are unmatched openings

def rough_brace_imbalance(text: str) -> int:
    total = 0
    for raw in text.splitlines():
        s = raw.split('%', 1)[0]
        total += s.count('{')
        total -= s.count('}')
    return total

def auto_fix_braces(text: str, how_many: int) -> str:
    if how_many <= 0:
        return text
    lines = text.splitlines()
    # insert before \end{document}
    end_doc_idx = None
    for idx, ln in enumerate(lines):
        if re.match(r'^\s*\\end\{document\}\s*$', ln):
            end_doc_idx = idx
            break
    if end_doc_idx is None:
        end_doc_idx = len(lines)
    lines[end_doc_idx:end_doc_idx] = ['}' for _ in range(how_many)]
    return '\n'.join(lines) + '\n'

# ---------------------- Pandoc helpers ----------------------

def find_pandoc(path: Optional[str]) -> str:
    if path:
        return path
    p = shutil.which("pandoc")
    if not p:
        print("ERROR: pandoc not found; install it or pass --pandoc PATH", file=sys.stderr)
        sys.exit(2)
    return p

def run_pandoc(pandoc: str, tex: str, out_docx: str, ref_docx: Optional[str]) -> None:
    cmd = [pandoc, "-f", "latex", "-t", "docx", tex, "-o", out_docx]
    if ref_docx:
        cmd.extend(["--reference-doc", ref_docx])
    subprocess.run(cmd, check=True)

# ---------------------- CLI ----------------------

def main():
    ap = argparse.ArgumentParser(description="Normalize LaTeX then optionally run Pandoc → DOCX.")
    ap.add_argument("--in", dest="inp", required=True, help="input .tex")
    ap.add_argument("--emit-tex", dest="emit_tex", default=None, help="write normalized .tex here (no pandoc)")
    ap.add_argument("--out", dest="out_docx", default=None, help="output .docx")
    ap.add_argument("--run-pandoc", action="store_true", help="actually run pandoc after normalizing")
    ap.add_argument("--pandoc", dest="pandoc_path", default=None, help="path to pandoc binary")
    ap.add_argument("--font-scheme", choices=["cmu","latin-modern"], default="cmu")
    ap.add_argument("--base-size", type=float, default=11.0)
    ap.add_argument("--debug-braces", action="store_true", help="print unmatched '{' line/col hotspots after normalization")
    ap.add_argument("--auto-fix-braces", action="store_true", help="if brace imbalance > 0, insert that many '}' before \\end{document}")
    args = ap.parse_args()

    raw = Path(args.inp).read_text(encoding="utf-8", errors="ignore")

    s = guard_glyphtounicode(raw)
    s = strip_macro_definitions(s)
    s = replace_invocations(s)
    s, stats = balance_envs(s)
    s = ensure_list_items(s)

    brace_delta = rough_brace_imbalance(s)
    hotspots = find_unmatched_open_braces(s) if brace_delta > 0 else []

    if args.auto_fix_braces and brace_delta > 0:
        s = auto_fix_braces(s, brace_delta)
        brace_delta = rough_brace_imbalance(s)

    norm_path = args.emit_tex or (Path(args.inp).with_suffix(".pandoc_ready.tex").as_posix())
    Path(norm_path).write_text(s, encoding="utf-8")

    print("Normalization report:")
    print(" - Missing env closers inserted:", stats["inserted_missing_ends"])
    print(" - Unmatched \\end{...} seen:", stats["unmatched_end_count"])
    print(" - Rough brace imbalance (should be 0):", brace_delta)
    print(" - Wrote normalized:", norm_path)

    if hotspots and args.debug_braces:
        print("\nUnmatched '{' hotspots (line:col):")
        for (ln, col, ctx) in hotspots[-10:]:  # show up to last 10
            print(f"  {ln}:{col}  {ctx.strip()}")

    if not args.run_pandoc:
        return

    if not args.out_docx:
        print("ERROR: --out is required with --run-pandoc", file=sys.stderr)
        sys.exit(2)

    fonts = CMU if args.font_scheme == "cmu" else LM
    ref_docx = make_reference_docx(fonts, base_size=args.base_size)

    try:
        pandoc = find_pandoc(args.pandoc_path)
        run_pandoc(pandoc, norm_path, args.out_docx, ref_docx)
        print("✓ DOCX written:", args.out_docx)
        print("✓ Used reference:", ref_docx)
    except subprocess.CalledProcessError as e:
        print("Pandoc failed with exit code", e.returncode, file=sys.stderr)
        cmd = [pandoc, "-f", "latex", "-t", "docx", norm_path, "-o", args.out_docx]
        if ref_docx:
            cmd += ["--reference-doc", ref_docx]
        print("Command:", " ".join(cmd), file=sys.stderr)
        sys.exit(e.returncode)

if __name__ == "__main__":
    main()
