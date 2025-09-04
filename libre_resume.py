#!/usr/bin/env python3
"""

Reads an input .docx resume, detects section headings and bullet/numbered lists,
and produces a neatly formatted resume .docx using python-docx.

Fixes:
- Removed use of private, non-existent run._r._add_fldChar API (which caused AttributeError).
- Adds optional --template to start from a .docx that already contains page-number fields
  in the footer/header. This is the recommended approach for dynamic page numbers with
  python-docx (fields are not supported directly).

Usage:
    pip install python-docx lxml
    python resume_reformatter_v2.py --in "/path/to/input.docx" --out "/path/to/output.docx"
    # Optional: use a template that already has PAGE/NUMPAGES fields in footer
    python resume_reformatter_v2.py --in input.docx --out output.docx --template base_with_page_numbers.docx
"""

import argparse
import sys
from typing import List, Dict, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECTION_TITLES = {
    "profile",
    "summary",
    "experience",
    "work experience",
    "projects",
    "skills",
    "technical skills",
    "education",
    "certifications",
    "publications",
    "awards",
}


def clean(text: str) -> str:
    return " ".join(text.split()).strip()


def detect_list_type(paragraph) -> Optional[str]:
    """
    Returns 'bullet', 'numbered', or None.

    Heuristic:
      1) Use paragraph style name if it includes 'List Bullet' or 'List Number'.
      2) Otherwise inspect XML for w:pPr/w:numPr and look up numFmt in the numbering part.
    """
    try:
        style_name = (paragraph.style.name or "").lower()
    except Exception:
        style_name = ""

    if "list bullet" in style_name or ("bullet" in style_name and "list" in style_name):
        return "bullet"
    if "list number" in style_name or "number" in style_name:
        return "numbered"

    # XML-based detection (w:numPr)
    try:
        pPr = paragraph._element.pPr  # type: ignore[attr-defined]
        numPr = getattr(pPr, "numPr", None)
    except Exception:
        numPr = None

    if numPr is None:
        return None

    try:
        numId_el = numPr.numId
        if numId_el is None:
            return None
        numId = numId_el.val
        numbering = paragraph.part.numbering_part.element  # CT_Numbering
        ns = numbering.nsmap

        # Find concrete <w:num w:numId="...">
        num_elems = numbering.xpath(f"w:num[@w:numId='{numId}']", namespaces=ns)
        if not num_elems:
            return "bullet"  # safe default

        # Resolve abstract numbering
        abstractNumId_elems = num_elems[0].xpath("w:abstractNumId", namespaces=ns)
        if not abstractNumId_elems:
            return "bullet"

        abstractNumId = abstractNumId_elems[0].get(qn("w:val"))
        # Look at level 0 numFmt as a proxy for type
        numFmt_elems = numbering.xpath(
            f"w:abstractNum[@w:abstractNumId='{abstractNumId}']/w:lvl[1]/w:numFmt",
            namespaces=ns,
        )
        if not numFmt_elems:
            return "bullet"

        fmt = numFmt_elems[0].get(qn("w:val")) or ""
        return "numbered" if fmt.lower() != "bullet" else "bullet"
    except Exception:
        # If anything goes sideways, assume bullet formatting
        return "bullet"


def is_probable_heading(paragraph_text: str, paragraph_style_name: str) -> bool:
    txt = clean(paragraph_text).strip(":").lower()
    if not txt:
        return False
    if (paragraph_style_name or "").lower().startswith("heading"):
        return True
    if txt in SECTION_TITLES:
        return True
    # Short, Title-case lines are often headings (e.g., "Experience")
    if len(txt) <= 32:
        words = [w for w in paragraph_text.split() if any(c.isalpha() for c in w)]
        if words and all(w[:1].isupper() for w in words):
            return True
    return False


def read_source(input_path: str) -> List[Dict]:
    """Parse the input DOCX into a simple list of paragraph records."""
    src = Document(input_path)
    records: List[Dict] = []
    for p in src.paragraphs:
        text = clean(p.text)
        if not text:
            continue
        try:
            style_name = p.style.name or ""
        except Exception:
            style_name = ""
        list_type = detect_list_type(p)
        is_list = list_type is not None
        is_heading = is_probable_heading(text, style_name)
        records.append(
            {
                "text": text,
                "style_name": style_name,
                "is_list": is_list,
                "list_type": list_type,  # 'bullet' | 'numbered' | None
                "is_heading": is_heading,
            }
        )
    return records


def build_output(records: List[Dict], output_path: str, template_path: Optional[str] = None) -> None:
    # If a template is provided, start from it (recommended if it contains page-number fields)
    doc = Document(template_path) if template_path else Document()

    # Page layout: compact professional margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base typography
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10.5)

    # Optional title block (best-effort from first non-empty text if looks like a name)
    possible_name = records[0]["text"] if records else None
    if possible_name and len(possible_name.split()) <= 6:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(possible_name)
        run.bold = True
        run.font.size = Pt(16)

    # Add body, grouping bullets under headings when found
    for rec in records:
        text = rec["text"]
        if rec["is_heading"]:
            doc.add_heading(text, level=1)
            continue

        if rec["is_list"]:
            style = "List Number" if rec["list_type"] == "numbered" else "List Bullet"
            para = doc.add_paragraph(text, style=style)
            para.paragraph_format.space_after = Pt(2)
        else:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)

    # IMPORTANT: No direct PAGE field insertion here, because python-docx doesn't expose a stable API
    # for fields. If you want automatic page numbers, create a template that has a page-number field
    # in the footer and pass it via --template. The field will remain intact and update in Word.

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Reformat a resume .docx with neat styles.")
    parser.add_argument("--in", dest="input_path", required=True, help="Path to input .docx resume")
    parser.add_argument("--out", dest="output_path", required=True, help="Path to write the formatted .docx")
    parser.add_argument("--template", dest="template_path", required=False, default=None,
                        help="Optional template .docx to start from (can contain footer page-number field)")
    args = parser.parse_args()

    records = read_source(args.input_path)
    if not records:
        print("No paragraphs found in the input document.", file=sys.stderr)
        sys.exit(2)

    build_output(records, args.output_path, template_path=args.template_path)
    print(f"✓ Wrote formatted resume to: {args.output_path}")


if __name__ == "__main__":
    main()
